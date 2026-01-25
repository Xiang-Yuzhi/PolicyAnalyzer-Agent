import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from typing import Dict, Any

class ReportGenerator:
    """
    负责将分析结果转换为标准 Word 文档
    应用易方达品牌格式标准
    """
    
    @staticmethod
    def set_efund_style(doc):
        """设置易方达品牌样式"""
        # 设置默认字体
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal'].font.size = Pt(10.5)  # 五号 = 10.5pt
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 设置英文字体
        doc.styles['Normal'].font.name = 'Times New Roman'
        
        # 段落格式
        doc.styles['Normal'].paragraph_format.line_spacing = Pt(18)
        doc.styles['Normal'].paragraph_format.space_before = Pt(6)  # 0.5行约6pt
        doc.styles['Normal'].paragraph_format.space_after = Pt(6)
        doc.styles['Normal'].paragraph_format.first_line_indent = Inches(0.25)  # 首行缩进2字符
        doc.styles['Normal'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 两端对齐

    @staticmethod
    def add_efund_heading(doc, text, level=1):
        """添加易方达风格的标题"""
        heading = doc.add_heading(text, level=level)
        
        # 设置标题颜色为易方达蓝
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 78, 157)
            
            if level == 1:
                run.font.name = '黑体'
                run.font.size = Pt(14)  # 4号 = 14pt
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            elif level == 2:
                run.font.name = '黑体'
                run.font.size = Pt(12)  # 小四 = 12pt
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
        
        return heading
    
    @staticmethod
    def generate_docx(analysis_data: Dict[str, Any], output_path: str = "report.docx") -> str:
        """
        输入: LLM 分析生成的 JSON (完整结构)
        输出: 生成文件的路径
        """
        doc = Document()
        
        # 应用易方达样式
        ReportGenerator.set_efund_style(doc)
        
        # 1. 标题
        heading = doc.add_heading('易方达基金 - 政策分析解读报告', 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 78, 157)
        
        # 2. 政策信息
        policy_info = analysis_data.get('selected_policy', {}) or {}
        policies_analyzed = analysis_data.get('policies_analyzed', [])
        
        if policy_info:
            policy_title = policy_info.get('title', '未命名政策')
            meta_text = f"发布机构: {policy_info.get('issuer', '-')}  发布时间: {policy_info.get('publish_date', '-')}  来源: {policy_info.get('url', '-')}"
        elif policies_analyzed:
            policy_title = f"组合政策分析报告 ({len(policies_analyzed)} 份)"
            meta_text = "分析对象: " + "、".join(policies_analyzed)
        else:
            policy_title = "未命名深度分析报告"
            meta_text = "投研深度政策分析"

        subtitle = doc.add_paragraph(policy_title)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].bold = True
        subtitle.runs[0].font.size = Pt(16)
        subtitle.paragraph_format.first_line_indent = Inches(0)
        
        # 元信息
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_para.paragraph_format.first_line_indent = Inches(0)
        meta_run = meta_para.add_run(meta_text)
        meta_run.italic = True
        meta_run.font.size = Pt(10)
        
        doc.add_paragraph()  # 空行

        
        # 3. 正文内容 (自适应所有返回的章节)
        content_data = analysis_data.get('docx_content', {}) or {}
        
        for section_title, paragraphs in content_data.items():
            if not paragraphs:
                continue
            
            # 添加易方达风格标题
            ReportGenerator.add_efund_heading(doc, section_title, level=1)
            
            # 添加段落内容
            if isinstance(paragraphs, list):
                for para_text in paragraphs:
                    if isinstance(para_text, str) and para_text.strip():
                        p = doc.add_paragraph(para_text)
                        # 应用正文格式
                        p.paragraph_format.first_line_indent = Inches(0.25)
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p.paragraph_format.line_spacing = Pt(18)
                        p.paragraph_format.space_before = Pt(6)
                        p.paragraph_format.space_after = Pt(6)
            else:
                p = doc.add_paragraph(str(paragraphs))
                p.paragraph_format.first_line_indent = Inches(0.25)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        
        # 4. 免责声明
        doc.add_page_break()

        footer = doc.add_paragraph("本报告仅供易方达内部投研参考，不构成公开投资建议。")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.paragraph_format.first_line_indent = Inches(0)
        for run in footer.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)
        
        # 保存
        try:
            doc.save(output_path)
        except PermissionError:
            output_path = f"report_{os.urandom(2).hex()}.docx"
            doc.save(output_path)
        
        return output_path